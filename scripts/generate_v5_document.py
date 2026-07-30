#!/usr/bin/env python3
"""Generate Noutq V5 DOCX from the project source-of-truth export."""

from __future__ import annotations

import argparse
import json
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlparse

from dotenv import load_dotenv
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


GREEN = "0F3D2E"
MID_GREEN = "1F6F54"
PALE_GREEN = "E8F4EF"
PALE_GOLD = "FFF7DA"
INK = "1F1F1F"
ARMENIAN_INK = "374151"
MUTED = "6B7280"
WHITE = "FFFFFF"
BORDER = "B8C9C1"
ARABIC_RE = re.compile(r"[\u0600-\u06ff]")
ARMENIAN_RE = re.compile(r"[\u0530-\u058f]")


UNIT_DESCRIPTIONS = {
    "C01": ("التعرّف إلى الحروف والحركات والتمييز السمعي قبل القراءة.", "Տառերը և շարժումները ճանաչել լսողական տարբերակումից սկսելով։"),
    "C02": ("الانتقال المتدرّج من شكل الحرف إلى العبارة القصيرة.", "Տառի ձևից աստիճանաբար անցնել կարճ արտահայտության։"),
    "C03": ("بدء محادثة والتعريف بالنفس بصيغتي المذكر والمؤنث.", "Սկսել զրույց և ներկայանալ՝ արական ու իգական դիմելաձևերով։"),
    "C04": ("فهم الأعداد من 1 إلى 10 واستخدامها في سياق.", "Հասկանալ և համատեքստում կիրառել 1–10 թվերը։"),
    "C05": ("السؤال عن الوقت والإجابة ثم وصف الطقس.", "Հարցնել ժամը, պատասխանել և նկարագրել եղանակը։"),
    "C06": ("تسمية أفراد الأسرة وتقديم معلومات قصيرة عنهم.", "Անվանել ընտանիքի անդամներին և կարճ ներկայացնել նրանց։"),
    "C07": ("طلب الطعام والشراب والتعبير عن الحاجة بأدب.", "Քաղաքավարի խնդրել ուտելիք և խմիչք։"),
    "C08": ("وصف المكان بحروف الجر وظروف المكان.", "Նկարագրել տեղը նախդիրներով և տեղի մակբայներով։"),
    "C09": ("التواصل في المدرسة والسؤال عن المادة المفضلة.", "Դպրոցում հաղորդակցվել և հարցնել սիրելի առարկայի մասին։"),
    "C10": ("السؤال عن الثمن وإتمام حوار شراء قصير.", "Հարցնել գինը և վարել կարճ գնման երկխոսություն։"),
    "C11": ("السؤال عن الطريق وإعطاء تعليمات للمذكر والمؤنث.", "Հարցնել ճանապարհը և ուղղություն տալ արական ու իգական ձևերով։"),
    "C12": ("وصف الروتين اليومي في جمل مترابطة.", "Նկարագրել առօրյան կապակցված նախադասություններով։"),
    "E01": ("امتداد اختياري للتعبير الطبيعي عن ألم بسيط.", "Ընտրովի բաժին՝ ցավը բնական ձևով արտահայտելու համար։"),
    "G01": ("تمييز التنوين وتوظيفه بعد بناء أساس تواصلي.", "Տարբերել թանվինը հաղորդակցական հիմքից հետո։"),
    "G02": ("مقدمة مبسّطة في الرفع والنصب والجر.", "Պարզ ներածություն ուղղական, հայցական և սեռական հոլովներին։"),
    "G03": ("تطبيق المطابقة في الجنس والعدد.", "Կիրառել սեռի և թվի համաձայնությունը։"),
    "G04": ("تمييز المذكر والمؤنث وصيغ الخطاب.", "Տարբերել արական, իգական և դիմելաձևերը։"),
    "G05": ("قراءة الحركات الإعرابية في سياقات قصيرة.", "Կարդալ իրաբական վերջավորությունները կարճ համատեքստերում։"),
    "R01": ("مراجعة متباعدة بالسياق بدل تكرار السؤال نفسه.", "Համատեքստային կրկնություն՝ նույն հարցը կրկնելու փոխարեն։"),
}

TYPE_AR = {
    "listen": "استماع",
    "listening": "استماع",
    "listening-discrimination": "تمييز سمعي",
    "speak": "نطق",
    "speaking": "إنتاج شفهي",
    "quiz": "اختيار",
    "contextual-choice": "اختيار سياقي",
    "match": "مطابقة",
    "matching": "مطابقة",
    "write": "كتابة",
    "writing": "كتابة",
    "reading": "قراءة",
    "production": "إنتاج",
    "mini-dialogue": "حوار",
    "sentence-completion": "إكمال",
    "classification": "تصنيف",
    "ordering": "ترتيب",
    "transformation": "تحويل",
    "error-correction": "تصحيح",
}

ACTIVITY_ICON = {
    "listen": "🎧",
    "listening": "🎧",
    "listening-discrimination": "🎧",
    "speak": "🗣️",
    "speaking": "🗣️",
    "quiz": "❓",
    "contextual-choice": "❓",
    "match": "🔗",
    "matching": "🔗",
    "write": "✍️",
    "writing": "✍️",
    "reading": "📖",
    "production": "🗣️",
    "mini-dialogue": "💬",
    "sentence-completion": "✏️",
    "classification": "🗂️",
    "ordering": "🔢",
    "transformation": "🔄",
    "error-correction": "🛠️",
}

# Accent visuel par piste — pas une nouvelle taxonomie, juste une couleur cohérente
# avec la palette de marque pour distinguer d'un coup d'œil le tronc commun (core),
# l'extension grammaticale optionnelle, la révision et l'évaluation.
TRACK_ACCENT = {
    "core": MID_GREEN,
    "grammar": "1D4E89",
    "optional": "B45309",
    "review": "6B4FA0",
    "assessment": "9B2C2C",
}


@dataclass
class QrCell:
    """Marqueur détecté par add_table() pour insérer une image QR au lieu
    d'un texte de statut. La décision d'éligibilité (includeInBookQr) n'est
    jamais reprise ici : elle vient déjà tranchée du manifeste audio."""
    audio_id: str
    qr_path: Path


def load_qr_metadata(path: Path) -> dict[str, Path]:
    """Absent (INCLUDE_AUDIO_QR=false ou pas encore généré) -> dict vide,
    jamais une erreur : le livre se génère alors sans QR, comme avant cette
    fonctionnalité."""
    if not path.exists():
        return {}
    payload = json.loads(path.read_text(encoding="utf-8"))
    # qrFile est déjà relatif à la racine du projet (voir generateAudioQr.ts),
    # donc résolu depuis le cwd — ce script est toujours lancé depuis la racine du repo.
    return {item["audioId"]: (Path.cwd() / item["qrFile"]).resolve() for item in payload["items"]}


def set_picture_name(shape, name: str, descr: str) -> None:
    """python-docx nomme l'image insérée d'après le nom de fichier source
    (ex: u1.4.png) dans un attribut XML interne, non lu par un lecteur normal
    mais visible en inspectant le document — on le remplace pour ne laisser
    aucun identifiant technique nulle part, même hors du texte affiché."""
    cnv_pr = shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr
    cnv_pr.set("name", name)
    cnv_pr.set("descr", descr)


def add_qr_cell(cell, audio_id: str, qr_path: Path, *, available: bool) -> None:
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # 2.0cm : dans la fourchette 18–22mm demandée, net à 300 DPI (voir generateAudioQr.ts).
    shape = run.add_picture(str(qr_path), width=Cm(2.0))
    set_picture_name(shape, "QR", "رمز QR")
    caption = cell.add_paragraph()
    caption.paragraph_format.space_after = Pt(0)
    caption.paragraph_format.space_before = Pt(0)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run("🔊 استمع / Լսիր" if available else "🔊 استمع (قريبًا)")
    set_run_font(caption_run, "Arial", 6.5, MUTED if available else "B45309")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=55, start=105, bottom=55, end=105) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_no_split(row, repeat_header=False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    if repeat_header:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        borders.append(el)


def set_table_left_accent(table, color: str, size: str = "24") -> None:
    """Épaissit et colore uniquement la bordure gauche (déjà posée par
    set_table_borders) — le liseré coloré qui distingue une piste (core/
    grammaire/révision/évaluation) d'un coup d'œil sur les cartes d'exercice."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = int(sum(widths_inches) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            twips = int(widths_inches[index] * 1440)
            cell.width = Inches(widths_inches[index])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(twips))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_bidi(paragraph, rtl: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        p_pr.append(bidi)
    elif not rtl and bidi is not None:
        p_pr.remove(bidi)


def set_run_font(run, family: str, size: float, color: str, bold=False, rtl=False) -> None:
    run.font.name = family
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), family)
    if rtl:
        rtl_el = r_pr.find(qn("w:rtl"))
        if rtl_el is None:
            rtl_el = OxmlElement("w:rtl")
            rtl_el.set(qn("w:val"), "1")
            r_pr.append(rtl_el)


def add_text(paragraph, text: str, size=10, bold=False, color=INK, rtl=None):
    if rtl is None:
        rtl = bool(ARABIC_RE.search(text)) and not bool(ARMENIAN_RE.search(text))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_bidi(paragraph, rtl)
    run = paragraph.add_run(text)
    family = "Arial" if rtl or not ARMENIAN_RE.search(text) else "Sylfaen"
    set_run_font(run, family, size, color, bold=bold, rtl=rtl)
    return run


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_mixed_cell(cell, text: str, *, header=False, size=9) -> None:
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    rtl = bool(ARABIC_RE.search(text)) and not bool(ARMENIAN_RE.search(text))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else (WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT)
    set_paragraph_bidi(p, rtl)
    run = p.add_run(text)
    family = "Arial" if rtl or header or not ARMENIAN_RE.search(text) else "Sylfaen"
    set_run_font(run, family, size, WHITE if header else (INK if rtl else ARMENIAN_INK), bold=header, rtl=rtl)


def add_heading(doc: Document, ar: str, hy: str, level=1, code: str | None = None) -> None:
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = level == 1
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_bidi(p, True)
    label = f"{code}  ·  {ar}" if code else ar
    run = p.add_run(label)
    set_run_font(run, "Arial", 16 if level == 1 else 12, GREEN, bold=True, rtl=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.keep_with_next = True
    p2.paragraph_format.space_after = Pt(6)
    add_text(p2, hy, size=10.5, bold=True, color=MID_GREEN, rtl=False)


def add_callout(doc: Document, ar: str, hy: str, fill=PALE_GREEN) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.58])
    set_table_borders(table, color=fill, size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    add_text(p, ar, size=10, color=GREEN, rtl=True)
    p2 = cell.add_paragraph()
    add_text(p2, hy, size=9, color=ARMENIAN_INK, rtl=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_unit_lead(doc: Document, ar: str, hy: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    shade_paragraph(p, PALE_GREEN)
    add_text(p, f"  {ar}  ", size=10, color=GREEN, rtl=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(7)
    shade_paragraph(p2, PALE_GREEN)
    add_text(p2, f"  {hy}  ", size=9, color=ARMENIAN_INK, rtl=False)


def add_table(doc: Document, headers: list[str], rows: Iterable[list[str]], widths: list[float], font_size=8.5, header_color=GREEN):
    row_values_list = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    # Mark every table's first row semantically as a header. Word only repeats it
    # when a table flows to another page, so this is safe for short tables too.
    set_row_no_split(table.rows[0], repeat_header=True)
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], header_color)
        add_mixed_cell(table.rows[0].cells[i], header, header=True, size=8.5)
    for row_values in row_values_list:
        row = table.add_row()
        set_row_no_split(row)
        for i, value in enumerate(row_values):
            if len(table.rows) % 2 == 0:
                set_cell_shading(row.cells[i], "F7FAF8")
            if isinstance(value, QrCell):
                add_qr_cell(row.cells[i], value.audio_id, value.qr_path, available=True)
            else:
                add_mixed_cell(row.cells[i], str(value), size=font_size)
    return table


def resolve_audio_status_cell(
    exercise_id: str,
    audio: dict[str, Any],
    qr_files: dict[str, Path],
    qr_warnings: list[str],
    fallback_labels: dict[str, str] | None = None,
) -> Any:
    """Statut audio unique, calculé une seule fois à partir du manifeste live —
    jamais du champ `audio` potentiellement obsolète embarqué dans curriculum.json.
    Retourne un QrCell (image) quand un QR existe pour cet id, sinon le texte de statut."""
    available = audio.get("status") == "available"
    if fallback_labels is not None:
        fallback = fallback_labels.get(audio.get("fallback"), audio.get("fallback", "—"))
        status = "متاح" if available else f"بديل: {fallback}"
    else:
        status = "متاح" if available else "غير متاح؛ تابع القراءة"
    if audio.get("includeInBookQr"):
        qr_path = qr_files.get(exercise_id)
        if qr_path is not None:
            return QrCell(exercise_id, qr_path)
        qr_warnings.append(f"{exercise_id}: مؤهّل لرمز QR لكنه غير مولَّد بعد (status={'available' if available else 'missing'})")
    return status


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    set_run_font(run, "Arial", 8, MUTED)


def configure_document(doc: Document, release: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.widow_control = True
    for name, size in (("Title", 26), ("Subtitle", 13), ("Heading 1", 16), ("Heading 2", 12), ("Heading 3", 10.5)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(GREEN)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(4)

    footer = section.footer
    p = footer.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Noutq — {release}  ·  نسخة المتعلّم / Սովորողի տարբերակ  ·  ")
    set_run_font(run, "Arial", 8, MUTED)
    add_page_number(p)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


FALLBACK_LABELS_AR = {
    "reading": "قراءة",
    "role-play": "تمثيل حوار",
    "writing": "كتابة",
    "teacher-read": "قراءة المعلّم",
}


def lesson_rows(
    data: dict[str, Any],
    unit: dict[str, Any],
    qr_files: dict[str, Path],
    qr_warnings: list[str],
) -> list[dict[str, Any]]:
    manifest = data["audioManifest"]["entries"]
    items: list[dict[str, Any]] = []
    for source_id in unit["legacySources"]:
        lesson = data["lessons"][source_id]
        for step in lesson["steps"]:
            exercise_id = f"{source_id}.{step['id']}"
            audio = manifest.get(exercise_id, {})
            status_cell = resolve_audio_status_cell(exercise_id, audio, qr_files, qr_warnings)
            items.append({
                "id": exercise_id,
                "type": step["type"],
                "typeLabel": TYPE_AR.get(step["type"], step["type"]),
                "arabic": step.get("arabic", ""),
                "translit": step.get("transliteration", "—") or "—",
                "armenian": step.get("armenian", ""),
                "statusCell": status_cell,
            })
    return items


def activity_rows(
    data: dict[str, Any],
    unit_id: str,
    qr_files: dict[str, Path],
    qr_warnings: list[str],
) -> list[dict[str, Any]]:
    manifest = data["audioManifest"]["entries"]
    items: list[dict[str, Any]] = []
    for activity in data["curriculum"]["newActivities"]:
        if activity["unit"] != unit_id:
            continue
        # Statut live, pas activity.get("audio", {}) qui peut être obsolète dans curriculum.json.
        audio = manifest.get(activity["id"], {})
        status_cell = resolve_audio_status_cell(activity["id"], audio, qr_files, qr_warnings, FALLBACK_LABELS_AR)
        items.append({
            "id": activity["id"],
            "type": activity["type"],
            "typeLabel": TYPE_AR.get(activity["type"], activity["type"]),
            "arabic": activity.get("arabic", ""),
            "translit": activity.get("transliteration", "—") or "—",
            "armenian": activity.get("armenian", ""),
            "statusCell": status_cell,
        })
    return items


PRACTICE_TYPES_DOUBLE = {"writing", "production", "mini-dialogue", "speak", "speaking"}


def add_practice_line(container) -> None:
    """Ligne réglée (bordure basse) pour recopier le mot/la phrase à la main —
    l'élément qui transforme les tableaux de référence en carnet d'exercices."""
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_activity_card(
    doc: Document,
    item: dict[str, Any],
    show_id: bool = True,
    show_missing_status: bool = True,
    accent_color: str = BORDER,
) -> None:
    """Bloc d'exercice façon carnet : texte + transcription + traduction +
    ligne(s) de recopiage + case « أتقنته » — remplace la ligne de tableau
    compacte pour que le livre se pratique dans l'ordre, pas seulement s'y réfère.
    show_id/show_missing_status à False pour le profil « manuel » grand public :
    ni identifiant technique (u1.4, v5.c02.01…), ni texte d'indisponibilité
    interne — un lecteur non technique ne doit voir que ce qui l'aide à apprendre."""
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.58])
    set_table_borders(table, color=BORDER, size="4")
    set_table_left_accent(table, accent_color)
    set_row_no_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=90, start=140, bottom=70, end=140)

    header = cell.paragraphs[0]
    clear_paragraph(header)
    header.paragraph_format.space_after = Pt(2)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_bidi(header, True)
    icon = ACTIVITY_ICON.get(item.get("type"), "")
    label = f"{icon} {item['typeLabel']}".strip() if icon else item["typeLabel"]
    meta_text = f"{item['id']}  ·  {label}" if show_id else label
    meta_run = header.add_run(meta_text)
    set_run_font(meta_run, "Arial", 7.5, MUTED, rtl=True)
    status_cell = item["statusCell"]
    if isinstance(status_cell, QrCell):
        header.add_run("   ")
        pic_run = header.add_run()
        shape = pic_run.add_picture(str(status_cell.qr_path), width=Cm(1.3))
        set_picture_name(shape, "QR", "رمز QR")
    elif status_cell and show_missing_status:
        status_run = header.add_run(f"   —   {status_cell}")
        set_run_font(status_run, "Arial", 7.5, MUTED, rtl=True)

    if item.get("arabic"):
        p_ar = cell.add_paragraph()
        p_ar.paragraph_format.space_before = Pt(3)
        p_ar.paragraph_format.space_after = Pt(1)
        add_text(p_ar, item["arabic"], size=15, bold=True, color=INK, rtl=True)

    if item.get("translit") and item["translit"] != "—":
        p_tr = cell.add_paragraph()
        p_tr.paragraph_format.space_after = Pt(1)
        p_tr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_tr.add_run(item["translit"])
        set_run_font(run, "Arial", 9, MUTED, rtl=False)
        run.font.italic = True

    if item.get("armenian"):
        p_hy = cell.add_paragraph()
        p_hy.paragraph_format.space_after = Pt(2)
        add_text(p_hy, item["armenian"], size=9.5, color=ARMENIAN_INK, rtl=False)

    if item.get("arabic"):
        for _ in range(2 if item.get("type") in PRACTICE_TYPES_DOUBLE else 1):
            add_practice_line(cell)

    footer = cell.add_paragraph()
    footer.paragraph_format.space_before = Pt(2)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_bidi(footer, True)
    box_run = footer.add_run("☐ أتقنته")
    set_run_font(box_run, "Arial", 9, GREEN, bold=True, rtl=True)
    hy_run = footer.add_run("   /   Տիրապետեցի")
    set_run_font(hy_run, "Arial", 8, MID_GREEN, rtl=False)

    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_unit_start_qr(
    doc: Document,
    unit_id: str,
    ordered_items: list[dict[str, Any]],
    qr_files: dict[str, Path],
    qr_warnings: list[str],
    accent_color: str = MID_GREEN,
) -> None:
    """QR unique en tête d'unité qui ouvre directement le premier pas de cette
    unité dans l'app (réutilise le même resolver /a/{id} et le même
    findStepByAudioId côté client — aucun changement app nécessaire)."""
    if not ordered_items:
        return
    first_id = ordered_items[0]["id"]
    qr_path = qr_files.get(first_id)
    if qr_path is None:
        qr_warnings.append(f"{first_id}: id de démarrage de l'unité {unit_id} sans QR (ajouter un override includeInBookQr si souhaité)")
        return
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.58])
    set_table_borders(table, color=accent_color, size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GREEN)
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.paragraph_format.space_before = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(qr_path), width=Cm(2.4))
    set_picture_name(shape, "QR", "امسح لبدء الوحدة")
    cap1 = cell.add_paragraph()
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(cap1, "امسح لبدء هذه الوحدة داخل التطبيق", size=9.5, bold=True, color=accent_color, rtl=True)
    cap2 = cell.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2.paragraph_format.space_after = Pt(4)
    add_text(cap2, "Սկանավորեք՝ այս բաժինը հավելվածում սկսելու համար", size=8.5, color=ARMENIAN_INK, rtl=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc: Document, data: dict[str, Any], show_stats: bool = True, tagline: tuple[str, str] | None = None) -> None:
    release = data["curriculum"]["release"]
    build_date = data["curriculum"]["buildDate"]

    # Bandeau plein de couleur en tête de couverture — plus proche d'une vraie
    # couverture de livre que du bloc de texte centré sur fond blanc d'avant.
    band = doc.add_table(rows=1, cols=1)
    set_table_geometry(band, [6.58])
    row = band.rows[0]
    row.height = Cm(9.5)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = band.cell(0, 0)
    set_cell_shading(cell, GREEN)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=200, start=200, bottom=200, end=200)

    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("✦  NOUTQ  ✦")
    set_run_font(run, "Arial", 34, WHITE, bold=True)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_bidi(p, True)
    run = p.add_run("منهج العربية للناطقين بالأرمنية")
    set_run_font(run, "Arial", 18, WHITE, bold=True, rtl=True)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Հայախոսների համար արաբերենի ուսումնական ծրագիր")
    set_run_font(run, "Arial", 13, WHITE, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    if show_stats:
        add_callout(
            doc,
            f"{release}: مسار أساسي مستقل، امتداد نحوي اختياري، معرّفات V4 محفوظة.",
            "Թողարկման թեկնածու․ անկախ հիմնական ուղի, ընտրովի քերականական ընդլայնում և պահպանված V4 նույնացուցիչներ։",
        )
        legacy_count = sum(len(lesson["steps"]) for lesson in data["lessons"].values())
        new_count = len(data["curriculum"]["newActivities"])
        exam_count = sum(len(a["items"]) for a in data["curriculum"]["assessments"])
        add_table(
            doc,
            ["Legacy محفوظ", "أنشطة V5 جديدة", "بنود التقييم"],
            [[str(legacy_count), str(new_count), str(exam_count)]],
            [2.18, 2.2, 2.2],
            font_size=11,
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"{release} · {build_date}", size=10, bold=True, color=MID_GREEN, rtl=False)
    elif tagline:
        ar, hy = tagline
        add_callout(doc, ar, hy)
    doc.add_page_break()


def add_contents(doc: Document, data: dict[str, Any], show_technical_columns: bool = True) -> None:
    add_heading(doc, "الفهرس والمسارات", "Բովանդակություն և ուղիներ", level=1)
    units = data["curriculum"]["units"]
    if show_technical_columns:
        rows = [
            [
                unit["id"],
                unit["titleAr"],
                unit["titleHy"],
                " + ".join(unit["legacySources"]) if unit["legacySources"] else "جديد",
                unit["track"],
            ]
            for unit in units
        ]
        add_table(
            doc,
            ["الرمز", "العربية", "Հայերեն", "مصدر V4", "المسار"],
            rows,
            [0.65, 1.85, 2.0, 1.0, 1.08],
            font_size=8.5,
        )
    else:
        rows = [[str(i + 1), unit["titleAr"], unit["titleHy"]] for i, unit in enumerate(units)]
        add_table(
            doc,
            ["#", "العربية", "Հայերեն"],
            rows,
            [0.5, 3.04, 3.04],
            font_size=9,
        )
    add_callout(
        doc,
        "يمكن إتمام المسار الأساسي دون امتداد القواعد الاختياري. يُعرض تقدّمك في كل مسار بصورة مستقلة داخل التطبيق.",
        "Հիմնական ուղին կարելի է ավարտել առանց ընտրովի քերականական ընդլայնման։ Ձեր առաջընթացը յուրաքանչյուր ուղում ցուցադրվում է հավելվածում առանձին։",
        fill=PALE_GOLD,
    )


def add_learning_guide(doc: Document, data: dict[str, Any]) -> None:
    add_heading(doc, "دليل التعلّم", "Ուսուցման ուղեցույց", level=1)
    add_table(
        doc,
        ["المرحلة", "العمل", "الزمن"],
        [
            ["1", "استرجاع من الذاكرة قبل فتح الدرس", "3 دقائق"],
            ["2", "استماع وقراءة موجّهة", "5 دقائق"],
            ["3", "إنتاج شفهي أو كتابي", "5 دقائق"],
            ["4", "تصحيح خطأين وتحديد موعد المراجعة", "دقيقتان"],
        ],
        [0.65, 4.65, 1.28],
        font_size=9,
    )
    add_callout(
        doc,
        "التكرار المتباعد: اليوم 0 / +1 / +3 / +7 / +14 / +30. معيار الإتقان 80٪ مع إنتاج صحيح دون نقحرة.",
        "Տարածված կրկնություն՝ 0 / +1 / +3 / +7 / +14 / +30։ Յուրացման շեմը՝ 80% և ինքնուրույն ճիշտ արտադրություն։",
    )

    add_heading(doc, "دليل النطق العربي–الأرمني", "Արաբերեն–հայերեն արտասանության ուղեցույց", level=2)
    guide = data["curriculum"]["pronunciationGuide"]
    add_callout(doc, guide["principleAr"], guide["principleHy"])
    rows = []
    for category in guide["categories"]:
        rows.append([
            category["id"],
            category["titleAr"],
            "، ".join(category["sounds"]),
            category["noteAr"],
            category["noteHy"],
        ])
    add_table(
        doc,
        ["الفئة", "التصنيف", "الأصوات", "طريقة التدريب", "Հայերեն"],
        rows,
        [0.45, 1.35, 1.28, 1.9, 1.6],
        font_size=8.2,
    )
    pairs = "  ·  ".join(guide["minimalPairs"])
    add_callout(
        doc,
        f"اختبار التمييز: {pairs}. سجّل ثم حقّق 8 إجابات صحيحة من 10.",
        "Տարբերակման ստուգում․ ձայնագրիր զույգերը և հասիր 10-ից 8 ճիշտ պատասխանի։",
        fill=PALE_GOLD,
    )


def add_integration(doc: Document, data: dict[str, Any], base_url: str | None) -> None:
    add_heading(doc, "حماية التكامل مع التطبيق", "Հավելվածի ինտեգրման պաշտպանություն", level=1, code="A02")
    add_callout(
        doc,
        "u1…u22 معرّفات Legacy ثابتة. C01…X02 طبقة تنظيم V5، والأنشطة الجديدة فقط تستعمل namespace v5.*.",
        "u1…u22-ը կայուն Legacy նույնացուցիչներ են։ C01…X02-ը V5 կազմակերպչական շերտն է, իսկ նոր վարժությունները՝ v5.*։",
    )
    rows = []
    for unit in data["curriculum"]["units"]:
        if unit["legacySources"]:
            rows.append([unit["id"], " + ".join(unit["legacySources"]), unit["titleAr"], "محفوظ"])
    add_table(doc, ["V5", "V4 Legacy", "المحتوى", "الحالة"], rows, [0.7, 1.3, 3.5, 1.08], font_size=8.5)

    add_heading(doc, "الصوت والروابط الثابتة", "Ձայն և կայուն հղումներ", level=2)
    available = sum(1 for e in data["audioManifest"]["entries"].values() if e["status"] == "available")
    missing = sum(1 for e in data["audioManifest"]["entries"].values() if e["status"] == "missing")
    qr_ar = (
        f"QR مفعّل على أساس {base_url}."
        if base_url
        else f"QR غير مضمّن في {data['curriculum']['release']} لأن الرابط العام canonical غير مضبوط؛ مسارات /a/{{audio_id}} جاهزة ومتحقّق منها."
    )
    qr_hy = (
        "QR-ը միացված է կայուն հանրային հասցեով։"
        if base_url
        else f"QR-ը {data['curriculum']['release']}-ում ներառված չէ, քանի որ հանրային canonical հասցեն սահմանված չէ։ /a/{{audio_id}} ուղիները պատրաստ են։"
    )
    add_table(
        doc,
        ["العنصر", "الحالة"],
        [
            ["audio_id", f"{len(data['audioManifest']['entries'])} mapping"],
            ["أصول محلية متاحة", str(available)],
            ["أصول ناقصة مع fallback", str(missing)],
            ["resolver", "/a/{audio_id}"],
            ["QR", qr_ar],
        ],
        [2.0, 4.58],
        font_size=9,
    )
    p = doc.add_paragraph()
    add_text(p, qr_hy, size=9, color=ARMENIAN_INK, rtl=False)

    add_heading(doc, "تغييرات النص المتوافقة", "Համատեղելի տեքստային փոփոխություններ", level=2)
    migration_rows = [
        [m["exerciseId"], m["from"], m["to"], "audio alias + SRS move"]
        for m in data["curriculum"]["textMigrations"]
    ]
    add_table(doc, ["المعرّف", "النص القديم", "النص المصحح", "الترحيل"], migration_rows, [0.85, 1.9, 2.45, 1.38], font_size=8)


def add_unit(
    doc: Document,
    data: dict[str, Any],
    unit: dict[str, Any],
    qr_files: dict[str, Path],
    qr_warnings: list[str],
    show_id: bool = True,
) -> None:
    unit_id = unit["id"]
    accent = TRACK_ACCENT.get(unit.get("track"), BORDER)
    add_heading(doc, unit["titleAr"], unit["titleHy"], level=1, code=unit_id if show_id else None)
    ar_desc, hy_desc = UNIT_DESCRIPTIONS.get(unit_id, ("", ""))
    if ar_desc:
        add_unit_lead(doc, ar_desc, hy_desc)

    items = lesson_rows(data, unit, qr_files, qr_warnings)
    new_items = activity_rows(data, unit_id, qr_files, qr_warnings)
    if unit_id in {"C02", "C03"}:
        items = new_items + items
        new_items = []

    add_unit_start_qr(doc, unit_id, items or new_items, qr_files, qr_warnings, accent_color=accent)

    if items:
        if unit_id in {"C02", "C03"}:
            add_heading(doc, "أنشطة الوحدة بترتيب V5", "Միավորի վարժությունները՝ V5 հերթականությամբ", level=2)
        else:
            add_heading(doc, "الأنشطة الموروثة المحفوظة", "Ժառանգված ու պահպանված վարժություններ", level=2)
        for item in items:
            add_activity_card(doc, item, show_id=show_id, show_missing_status=show_id, accent_color=accent)
    if new_items:
        add_heading(doc, "إضافات V5", "V5 հավելումներ", level=2)
        for item in new_items:
            add_activity_card(doc, item, show_id=show_id, show_missing_status=show_id, accent_color=accent)

    production = [
        a for a in data["curriculum"]["newActivities"]
        if a["unit"] == unit_id and a["type"] in {"mini-dialogue", "production", "writing"}
    ]
    if production and unit_id not in {"C03", "C05"}:
        activity = production[-1]
        add_callout(
            doc,
            f"مهمّة إنتاج: {activity['arabic']}",
            activity["armenian"],
            fill=PALE_GOLD,
        )


def add_review(
    doc: Document,
    data: dict[str, Any],
    unit: dict[str, Any],
    qr_files: dict[str, Path],
    qr_warnings: list[str],
    show_id: bool = True,
) -> None:
    accent = TRACK_ACCENT.get("review", BORDER)
    add_heading(doc, unit["titleAr"], unit["titleHy"], level=1, code="R01" if show_id else None)
    add_callout(
        doc,
        "المراجعة سياقية: لا يُعاد السؤال نفسه مباشرة، بل تُستعمل المهارة في مهمة جديدة.",
        "Կրկնությունը համատեքստային է․ նույն հարցը անմիջապես չի կրկնվում։",
    )
    legacy = lesson_rows(data, unit, qr_files, qr_warnings)
    add_unit_start_qr(doc, "R01", legacy, qr_files, qr_warnings, accent_color=accent)
    if legacy:
        add_heading(doc, "مرجع V4 المحفوظ", "Պահպանված V4 հղում", level=2)
        for item in legacy:
            add_activity_card(doc, item, show_id=show_id, show_missing_status=show_id, accent_color=accent)
    manifest = data["audioManifest"]["entries"]
    items = []
    for item in data["curriculum"]["review"]:
        # Statut live, pas item.get("audio", {}) qui peut être obsolète dans curriculum.json.
        audio = manifest.get(item["id"], {})
        status_cell = resolve_audio_status_cell(item["id"], audio, qr_files, qr_warnings, FALLBACK_LABELS_AR)
        items.append({
            "id": item["id"],
            "type": item["type"],
            "typeLabel": TYPE_AR.get(item["type"], item["type"]),
            "arabic": item["promptAr"],
            "translit": None,
            "armenian": item["promptHy"],
            "statusCell": status_cell,
        })
    add_heading(doc, "مراجعة V5 المتنوعة", "V5 բազմազան կրկնություն", level=2)
    for item in items:
        add_activity_card(doc, item, show_id=show_id, show_missing_status=show_id, accent_color=accent)


def add_assessment(
    doc: Document,
    data: dict[str, Any],
    assessment: dict[str, Any],
    qr_files: dict[str, Path],
    qr_warnings: list[str],
    show_id: bool = True,
) -> None:
    add_heading(doc, assessment["titleAr"], assessment["titleHy"], level=1, code=assessment["id"] if show_id else None)
    if assessment["id"] == "X01":
        add_callout(
            doc,
            "يقيس الاستماع والقراءة والمفردات والتواصل والكتابة والكلام. لا يعتمد على مفهوم نحوي اختياري.",
            "Գնահատվում են լսելը, կարդալը, բառապաշարը, հաղորդակցությունը, գրելը և խոսելը՝ առանց ընտրովի քերականության։",
        )
    else:
        add_callout(
            doc,
            "اختبار إضافي مستقل. لا يُشترط لإتمام المسار الأساسي.",
            "Առանձին լրացուցիչ քննություն է և պարտադիր չէ հիմնական ուղու համար։",
            fill=PALE_GOLD,
        )
    manifest = data["audioManifest"]["entries"]
    rows = []
    for i, item in enumerate(assessment["items"]):
        status_cell = resolve_audio_status_cell(item["id"], manifest.get(item["id"], {}), qr_files, qr_warnings)
        if not show_id and not isinstance(status_cell, QrCell):
            status_cell = ""  # pas de texte de statut technique dans le profil grand public
        rows.append([
            item["id"] if show_id else str(i + 1),
            TYPE_AR.get(item["type"], item["type"]),
            item["skill"],
            item["promptAr"],
            item["promptHy"],
            status_cell,
        ])
    add_table(
        doc,
        ["الرقم" if not show_id else "المعرّف", "النوع", "المهارة", "المهمة", "Հայերեն", "الصوت"],
        rows,
        [0.7, 0.65, 0.75, 1.98, 1.55, 0.95],
        font_size=8.1,
        header_color=TRACK_ACCENT["assessment"],
    )
    add_heading(doc, "سلّم مختصر", "Կարճ գնահատման սանդղակ", level=2)
    add_table(
        doc,
        ["المعيار", "4", "3", "2", "1–0"],
        [
            ["إنجاز المهمة", "كامل وواضح", "نقص بسيط", "نصف المطلوب", "غير منجز"],
            ["الدقة", "أخطاء قليلة", "أخطاء ملحوظة", "تعيق أحيانًا", "تعيق الفهم"],
            ["الاستقلال", "دون دعم", "دعم بسيط", "دعم متكرر", "غير مستقل"],
        ],
        [1.45, 1.28, 1.28, 1.28, 1.29],
        font_size=8.5,
    )


def add_glossary(doc: Document, data: dict[str, Any], show_id: bool = True) -> None:
    add_heading(doc, "المعجم المنظّم", "Կազմակերպված բառարան", level=1, code="A01" if show_id else None)
    add_callout(
        doc,
        "المعجم مصدره قائمة مفردات منسّقة؛ لا يضم تعليمات التمارين أو الأسئلة الناقصة.",
        "Բառարանը կազմված է խմբագրված բառացանկից և չի ներառում վարժությունների հրահանգներ կամ թերի հարցեր։",
    )
    add_heading(doc, "أ. الكلمات", "Ա. Բառեր", level=2)
    if show_id:
        word_rows = [
            [item["arabic"], item["transliteration"], item["armenian"], item["partOfSpeech"], item["relatedForm"] or "—", item["source"]]
            for item in data["glossary"]["words"]
        ]
        add_table(doc, ["الكلمة", "النطق", "Հայերեն", "النوع", "جمع/مؤنث", "المصدر"], word_rows, [1.1, 1.05, 1.55, 0.9, 1.2, 0.78], font_size=8.2)
    else:
        word_rows = [
            [item["arabic"], item["transliteration"], item["armenian"], item["partOfSpeech"], item["relatedForm"] or "—"]
            for item in data["glossary"]["words"]
        ]
        add_table(doc, ["الكلمة", "النطق", "Հայերեն", "النوع", "جمع/مؤنث"], word_rows, [1.25, 1.2, 1.78, 1.1, 1.25], font_size=8.4)
    add_heading(doc, "ب. التعبيرات والجمل", "Բ. Արտահայտություններ և նախադասություններ", level=2)
    if show_id:
        expression_rows = [
            [item["arabic"], item["transliteration"], item["armenian"], item["source"]]
            for item in data["glossary"]["expressions"]
        ]
        add_table(doc, ["التعبير", "النطق", "Հայերեն", "المصدر"], expression_rows, [2.1, 1.45, 2.15, 0.88], font_size=8.3)
    else:
        expression_rows = [
            [item["arabic"], item["transliteration"], item["armenian"]]
            for item in data["glossary"]["expressions"]
        ]
        add_table(doc, ["التعبير", "النطق", "Հայերեն"], expression_rows, [2.3, 1.6, 2.68], font_size=8.6)


def add_release_note(doc: Document, data: dict[str, Any], base_url: str | None) -> None:
    add_heading(doc, "ملاحظة الإصدار", "Թողարկման նշում", level=1)
    add_callout(
        doc,
        "هذه الوثيقة مولّدة من مصدر الحقيقة في المشروع. أي تعديل لاحق يجب أن يبدأ من البيانات ثم يعيد التوليد.",
        "Այս փաստաթուղթը ստեղծվել է նախագծի source of truth-ից։ Հետագա փոփոխությունները պետք է արվեն տվյալներում և վերագեներացվեն։",
    )
    add_table(
        doc,
        ["البوابة", "النتيجة"],
        [
            ["V4", "لم تُعدّل؛ u1…u22 محفوظة"],
            ["u8.71", "معرّف حقيقي ومتحقّق"],
            ["X01 / X02", "منفصلان؛ X02 اختياري"],
            ["الصوت", "كل نشاط استماع له mapping أو fallback"],
            ["QR", "مفعّل" if base_url else "غير مفعّل حتى ضبط canonical base URL"],
            ["المعجم", f"{len(data['glossary']['words'])} كلمة + {len(data['glossary']['expressions'])} تعبيرًا"],
            ["الإصدارات", f"schema {data['curriculum']['schemaVersion']} · migration {data['curriculum']['migrationVersion']} · audio {data['curriculum']['audioManifestVersion']}"],
        ],
        [2.0, 4.58],
        font_size=9,
    )


def add_foreword_manuel(doc: Document) -> None:
    """Mot d'ouverture + mention des auteurs — noms seuls, sans titre ni
    affiliation (aucune information non vérifiable sur les auteurs n'est
    inventée), au même format que la section correspondante ajoutée à
    l'article académique associé à ce projet."""
    add_heading(doc, "مقدّمة", "Ներածություն", level=1)
    p1 = doc.add_paragraph()
    add_text(
        p1,
        "تأتي هذه الوثيقة استجابةً لحاجة متنامية لدى الناطقين بالأرمنية إلى مورد تعليمي "
        "متكامل لتعلّم اللغة العربية، وهي لغة ذات امتداد حضاري وثقافي واسع ترتبط بها أرمينيا "
        "بروابط تاريخية وثقافية عميقة. ويقترح هذا الدليل، بمرافقة تطبيق Noutq، مسارًا تعليميًا "
        "تدريجيًا يراعي خصوصية المتعلّم الأرمني على المستويات الصوتية والكتابية والتواصلية، "
        "جامعًا بين الممارسة الورقية التقليدية والتطبيق الرقمي الحديث.",
        size=10.5,
        color=INK,
        rtl=True,
    )
    p1.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph()
    add_text(
        p2,
        "Այս փաստաթուղթը ստեղծվել է հայախոս սովորողների աճող կարիքին ի պատասխան՝ ունենալու "
        "արաբերենի ուսուցման ամբողջական միջոց։ Արաբերենը հարուստ քաղաքակրթական ու մշակութային "
        "ժառանգություն ունեցող լեզու է, որի հետ Հայաստանը կապված է խորը պատմական ու "
        "մշակութային կապերով։ Այս ուղեցույցը՝ Noutq հավելվածին զուգահեռ, առաջարկում է "
        "աստիճանական ուսումնական ուղի՝ հաշվի առնելով հայախոս սովորողի առանձնահատկությունները "
        "հնչյունաբանական, գրավոր և հաղորդակցական մակարդակներում։",
        size=9.5,
        color=ARMENIAN_INK,
        rtl=False,
    )
    p2.paragraph_format.space_after = Pt(12)

    p3 = doc.add_paragraph()
    add_text(p3, "المؤلّفون: خالد محمد أزلماض، محمد شوقي، أمين أمهان، صونا طونيكيان", size=10.5, bold=True, color=GREEN, rtl=True)
    p3.paragraph_format.space_after = Pt(14)


def add_about_manuel(doc: Document, data: dict[str, Any], has_qr: bool) -> None:
    """Page de présentation du profil « manuel » : ce que sont Noutq (l'app) et
    ce livre, comment les utiliser ensemble — en arabe/arménien uniquement,
    registre pédagogique, sans aucun vocabulaire d'ingénierie."""
    add_heading(doc, "عن التطبيق وهذا الدليل", "Հավելվածի և այս ուղեցույցի մասին", level=1)
    p1 = doc.add_paragraph()
    add_text(
        p1,
        "Noutq تطبيق لتعلّم اللغة العربية موجَّه للناطقين بالأرمنية، يجمع بين الاستماع "
        "والقراءة والكتابة والمحادثة ضمن مسار تدريجي يراعي حاجات المتعلّم الناطق بالأرمنية تحديدًا. "
        "يعتمد التطبيق على التكرار المتباعد لترسيخ المفردات والتراكيب في الذاكرة طويلة المدى، "
        "وعلى تصحيح فوري لكل محاولة نطق أو كتابة.",
        size=10.5,
        color=INK,
        rtl=True,
    )
    p1.paragraph_format.space_after = Pt(6)
    p2 = doc.add_paragraph()
    add_text(
        p2,
        "Noutq-ը արաբերենի ուսուցման հավելված է՝ ուղղված հայախոս սովորողներին։ Այն համատեղում է "
        "լսելը, կարդալը, գրելը և խոսակցական պրակտիկան աստիճանական ուսումնական ուղու շրջանակում՝ "
        "հաշվի առնելով հատկապես հայախոս սովորողի կարիքները։ Հավելվածն օգտագործում է տարածված "
        "կրկնության մեթոդը՝ բառապաշարն ու կառույցները երկարաժամկետ հիշողության մեջ ամրապնդելու համար։",
        size=9.5,
        color=ARMENIAN_INK,
        rtl=False,
    )
    p2.paragraph_format.space_after = Pt(10)

    add_heading(doc, "كيف تستعمل هذا الدليل مع التطبيق", "Ինչպես օգտագործել այս ուղեցույցը հավելվածի հետ", level=2)
    p3 = doc.add_paragraph()
    add_text(
        p3,
        "هذا الدليل مرافق للتطبيق لا بديل عنه: اقرأ الكلمة أو الجملة، انطقها بصوت مسموع، "
        "ثم اكتبها في المساحة المخصّصة تحتها لتثبيت شكلها. ضع علامة في المربّع "
        "«☐ أتقنته» بعد أن تكون قادرًا على قراءة النشاط وترجمته دون مساعدة.",
        size=10.5,
        color=INK,
        rtl=True,
    )
    p3.paragraph_format.space_after = Pt(6)
    p4 = doc.add_paragraph()
    add_text(
        p4,
        "Այս ուղեցույցը հավելվածին ուղեկցող է, ոչ թե փոխարինող. կարդացեք բառը կամ նախադասությունը, "
        "արտասանեք բարձրաձայն, ապա գրեք այն ներքևի հատուկ տողում՝ ձևը հիշողության մեջ ամրապնդելու համար։ "
        "Նշեք «☐ Տիրապետեցի» վանդակը այն բանից հետո, երբ կարողանաք կարդալ ու թարգմանել առաջադրանքն ինքնուրույն։",
        size=9.5,
        color=ARMENIAN_INK,
        rtl=False,
    )
    p4.paragraph_format.space_after = Pt(10)

    if has_qr:
        add_callout(
            doc,
            "كل وحدة تبدأ برمز QR: امسحه بكاميرا هاتفك ليفتح التطبيق مباشرة على أوّل نشاط "
            "في تلك الوحدة، فتستمع إلى النطق الصحيح وتتابع التمرين هناك.",
            "Յուրաքանչյուր բաժին սկսվում է QR կոդով. սկանավորեք այն ձեր հեռախոսի տեսախցիկով, "
            "որպեսզի հավելվածը բացվի անմիջապես տվյալ բաժնի առաջին առաջադրանքի վրա։",
        )


def build(
    data: dict[str, Any],
    output: Path,
    base_url: str | None,
    qr_files: dict[str, Path],
    qr_report: Path | None = None,
) -> None:
    doc = Document()
    qr_warnings: list[str] = []
    configure_document(doc, data["curriculum"]["release"])
    add_cover(doc, data)
    add_contents(doc, data)
    add_learning_guide(doc, data)
    add_integration(doc, data, base_url)

    units_by_id = {u["id"]: u for u in data["curriculum"]["units"]}
    for unit_id in ["C01", "C02", "C03", "C04", "C05", "C06", "E01", "C07", "C08", "C09", "C10", "C11", "C12"]:
        add_unit(doc, data, units_by_id[unit_id], qr_files, qr_warnings)
    for unit_id in ["G01", "G02", "G03", "G04", "G05"]:
        add_unit(doc, data, units_by_id[unit_id], qr_files, qr_warnings)
    add_review(doc, data, units_by_id["R01"], qr_files, qr_warnings)
    for assessment in data["curriculum"]["assessments"]:
        add_assessment(doc, data, assessment, qr_files, qr_warnings)
    add_glossary(doc, data)
    add_release_note(doc, data, base_url)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = f"Noutq {data['curriculum']['release']} — منهج العربية للناطقين بالأرمنية"
    doc.core_properties.subject = "Bilingual Arabic–Armenian A1 curriculum"
    doc.core_properties.author = "Noutq"
    doc.core_properties.keywords = "Arabic, Armenian, A1, Noutq, V5"
    doc.save(output)

    # Rapport interne uniquement : jamais imprimé dans le livre lui-même.
    if qr_report is not None:
        qr_report.parent.mkdir(parents=True, exist_ok=True)
        qr_report.write_text(
            json.dumps(
                {
                    "generatedAt": data["curriculum"].get("buildDate"),
                    "qrEligibleCount": sum(
                        1 for e in data["audioManifest"]["entries"].values() if e.get("includeInBookQr")
                    ),
                    "qrFilesLoaded": len(qr_files),
                    "warnings": sorted(set(qr_warnings)),
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        if qr_warnings:
            print(f"[qr] {len(set(qr_warnings))} élément(s) éligible(s) QR sans image générée — voir {qr_report}")


def build_manuel(
    data: dict[str, Any],
    output: Path,
    base_url: str | None,
    qr_files: dict[str, Path],
    qr_report: Path | None = None,
) -> None:
    """Profil « manuel » : document d'apprentissage grand public, arabe/arménien
    uniquement — aucun identifiant technique (u1.4, C01, schemaVersion…), aucune
    section d'intégration/infrastructure. Le contenu pédagogique et les QR sont
    les mêmes que le profil complet ; seule la présentation change."""
    doc = Document()
    qr_warnings: list[str] = []
    configure_document(doc, data["curriculum"]["release"])
    add_cover(
        doc,
        data,
        show_stats=False,
        tagline=(
            "الدليل الرسمي المرافق لتطبيق Noutq — منهج متكامل لتعلّم العربية موجَّه للناطقين بالأرمنية.",
            "Noutq հավելվածին ուղեկցող պաշտոնական ուղեցույցը՝ արաբերենի ուսուցման ամբողջական ծրագիր հայախոսների համար։",
        ),
    )
    add_foreword_manuel(doc)
    add_about_manuel(doc, data, has_qr=bool(base_url and qr_files))
    add_contents(doc, data, show_technical_columns=False)
    add_learning_guide(doc, data)

    units_by_id = {u["id"]: u for u in data["curriculum"]["units"]}
    for unit_id in ["C01", "C02", "C03", "C04", "C05", "C06", "E01", "C07", "C08", "C09", "C10", "C11", "C12"]:
        add_unit(doc, data, units_by_id[unit_id], qr_files, qr_warnings, show_id=False)
    for unit_id in ["G01", "G02", "G03", "G04", "G05"]:
        add_unit(doc, data, units_by_id[unit_id], qr_files, qr_warnings, show_id=False)
    add_review(doc, data, units_by_id["R01"], qr_files, qr_warnings, show_id=False)
    for assessment in data["curriculum"]["assessments"]:
        add_assessment(doc, data, assessment, qr_files, qr_warnings, show_id=False)
    add_glossary(doc, data, show_id=False)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Noutq — الدليل الرسمي / Պաշտոնական ուղեցույց"
    doc.core_properties.subject = "دليل تعلّم اللغة العربية للناطقين بالأرمنية"
    doc.core_properties.author = "Noutq"
    doc.core_properties.keywords = "Arabic, Armenian, Noutq"
    doc.save(output)

    if qr_report is not None:
        qr_report.parent.mkdir(parents=True, exist_ok=True)
        qr_report.write_text(
            json.dumps(
                {
                    "generatedAt": data["curriculum"].get("buildDate"),
                    "qrEligibleCount": sum(
                        1 for e in data["audioManifest"]["entries"].values() if e.get("includeInBookQr")
                    ),
                    "qrFilesLoaded": len(qr_files),
                    "warnings": sorted(set(qr_warnings)),
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        if qr_warnings:
            print(f"[qr] {len(set(qr_warnings))} élément(s) éligible(s) QR sans image générée — voir {qr_report}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--base-url", default=None)
    parser.add_argument(
        "--qr-metadata",
        type=Path,
        default=Path("output/qr/audio-qr-manifest.json"),
        help="Manifeste produit par scripts/generateAudioQr.ts. Absent -> livre généré sans QR (comportement inchangé).",
    )
    parser.add_argument(
        "--qr-report",
        type=Path,
        default=Path("reports/qr-eligible-missing.json"),
        help="Rapport interne (jamais dans le livre) des ids éligibles QR sans image générée.",
    )
    parser.add_argument(
        "--no-qr",
        action="store_true",
        help="Force la génération sans image QR même si NOUTQ_PUBLIC_BASE_URL est défini (ex: .env partagé).",
    )
    parser.add_argument(
        "--profile",
        choices=["full", "manuel"],
        default="full",
        help="'full' = document complet avec sections techniques (comportement historique). "
        "'manuel' = guide grand public arabe/arménien uniquement, sans identifiants techniques ni sections d'intégration.",
    )
    args = parser.parse_args()
    load_dotenv()  # même .env que les scripts TS ; n'écrase jamais une variable déjà présente dans l'environnement.
    data = json.loads(args.data.read_text(encoding="utf-8"))
    raw_base_url = args.base_url or os.environ.get("NOUTQ_PUBLIC_BASE_URL")
    base_url = None
    if raw_base_url:
        parsed = urlparse(raw_base_url.strip())
        if (
            parsed.scheme != "https"
            or not parsed.netloc
            or parsed.query
            or parsed.fragment
            or parsed.username
            or parsed.password
        ):
            raise ValueError("NOUTQ_PUBLIC_BASE_URL must be a clean HTTPS URL")
        base_url = raw_base_url.strip().rstrip("/")
    qr_files = load_qr_metadata(args.qr_metadata) if (base_url and not args.no_qr) else {}
    if base_url and not args.no_qr and not qr_files and args.qr_metadata.exists():
        print(f"[qr] {args.qr_metadata} lu mais vide — aucune image QR ne sera intégrée.")
    builder = build_manuel if args.profile == "manuel" else build
    builder(data, args.output, base_url, qr_files, qr_report=args.qr_report)
    print(f"Generated {args.output} (profile={args.profile})")


if __name__ == "__main__":
    main()
